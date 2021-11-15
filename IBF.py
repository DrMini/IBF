# -*- coding: utf-8 -*-
"""
Created on Sat May 15 16:49:41 2021

@authors 1: Dr. V K Mini, India Meteorological Department, Thiruvananthapuram
2: Lakshmi Jayalal, Research Scholar, Indian Institute of Technology, Madras

If you in any way use this code for research that results in publications,
please credit the authors.
"""
import ssl
import urllib
import urllib.request as urllib2
import glob
from PIL import Image
from datetime import date,datetime,timedelta
today = date.today()
yesterday=today-timedelta(days=1)
from pytz import timezone
import pandas as pd
import tkinter as tk
from tkinter import filedialog
import numpy as np
import array as arr

# SELECTION OF DISTRICTS TO ISSUE IBF --- 
ibfdist =arr.array('i',[0,0,0,0,0,0,0,0,0,0,0,0,0,0])

# Create an instance of tkinter frame
win = tk.Tk()
win.title("Select IBF Issued Districts : ")

# Set the geometry of Tkinter frame
win.geometry("350x450")

# Define Function to print the input value
def display_input():
   global ibfdist 
   ibfdist[0] = varibf[0].get()
   ibfdist[1] = varibf[1].get()
   ibfdist[2] = varibf[2].get()
   ibfdist[3] = varibf[3].get()
   ibfdist[4] = varibf[4].get()
   ibfdist[5] = varibf[5].get()
   ibfdist[6] = varibf[6].get()
   ibfdist[7] = varibf[7].get()
   ibfdist[8] = varibf[8].get()
   ibfdist[9] = varibf[9].get()
   ibfdist[10] = varibf[10].get()
   ibfdist[11] = varibf[11].get()
   ibfdist[12] = varibf[12].get()
   ibfdist[13] = varibf[13].get()
    
 
# Define empty variables
varibf = []
for i in range(14):
    varibf.append(tk.IntVar())

# Define a Checkbox
t1 = tk.Checkbutton(win, text="ALAPPUZHA", variable=varibf[0], onvalue=1, offvalue=0, command=display_input).grid(row = 0, column = 0, ipadx =2,  ipady = 2, sticky = tk.W)
#t1.pack()                                                        
t2 = tk.Checkbutton(win, text="ERNAKULAM", variable=varibf[1], onvalue=1, offvalue=0, command=display_input).grid(row = 1, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t2.pack()                                                        
t3 = tk.Checkbutton(win, text="IDUKKI", variable=varibf[2], onvalue=1, offvalue=0, command=display_input).grid(row = 2, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t3.pack()                                                     
t4 = tk.Checkbutton(win, text="KANNUR", variable=varibf[3], onvalue=1, offvalue=0, command=display_input).grid(row = 3, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t4.pack()                                                        
t5 = tk.Checkbutton(win, text="KASARAGOD", variable=varibf[4], onvalue=1, offvalue=0, command=display_input).grid(row = 4, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t5.pack()                                                       
t6 = tk.Checkbutton(win, text="KOLLAM", variable=varibf[5],  onvalue=1, offvalue=0, command=display_input).grid(row = 5, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t6.pack()                                                        
t7 = tk.Checkbutton(win, text="KOTTAYAM", variable=varibf[6], onvalue=1, offvalue=0, command=display_input).grid(row = 6, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t7.pack()                                                        
t8 = tk.Checkbutton(win, text="KOZHIKODE", variable=varibf[7], onvalue=1, offvalue=0, command=display_input).grid(row = 7, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t8.pack()                                                        
t9 = tk.Checkbutton(win, text="MALAPPURAM", variable=varibf[8], onvalue=1, offvalue=0, command=display_input).grid(row = 8, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t9.pack()                                                        
t10 = tk.Checkbutton(win, text="PALAKKAD", variable=varibf[9], onvalue=1, offvalue=0, command=display_input).grid(row = 9, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t10.pack()
t11 = tk.Checkbutton(win, text="PATHANAMTHITTA", variable=varibf[10], onvalue=1, offvalue=0, command=display_input).grid(row = 10, column = 0, ipadx =2, ipady = 2, sticky = tk.W) 
#t11.pack()                                                             
t12 = tk.Checkbutton(win, text="THIRUVANANTHAPURAM", variable=varibf[11], onvalue=1, offvalue=0, command=display_input).grid(row = 11, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t12.pack()
t13 = tk.Checkbutton(win, text="THRISSUR", variable=varibf[12], onvalue=1, offvalue=0, command=display_input).grid(row = 12, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t13.pack()                                                 
t14 = tk.Checkbutton(win, text="WAYANAD", variable=varibf[13],  onvalue=1, offvalue=0, command=display_input).grid(row = 13, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t14.pack()
t15 = tk.Button(win, text='Submit', command=win.destroy).grid(row = 14, column = 0, ipadx =2, ipady = 2, sticky = tk.W)
#t15.pack()

win.mainloop()

ibfyes = 0

#END OF CODE SELECTION OF IBF  DISTRICTS





root = tk.Tk()
root.withdraw()
ssl._create_default_https_context = ssl._create_unverified_context
# file_path = filedialog.askopenfilename()
format = "%H%M"
from docx.shared import Inches

def credentials(url, username="imd", password="Forecast#imd@2020" ):
    p = urllib2.HTTPPasswordMgrWithDefaultRealm()
    p.add_password(None, url, username, password)
    handler = urllib2.HTTPBasicAuthHandler(p)
    opener = urllib2.build_opener(handler)
    urllib2.install_opener(opener)
    
filenames=["https://www.satellite.imd.gov.in/imgr/asiasec_ir2.jpg","https://internal.imd.gov.in/section/dwr/img/caz_tvm.gif"\
            ,"https://internal.imd.gov.in/section/dwr/img/caz_koc.gif"]
Savefile=["LatestSatellitePicture.png","LatestRadarTVM.png","LatestRadarKOC.png"]
for file,sn in zip(filenames,Savefile):
    credentials(file)
    urllib.request.urlretrieve(file,sn)
from docx import Document
def replace_text_in_paragraph(paragraph, key, value):
    if key in paragraph.text:
        inline = paragraph.runs
        for item in inline:
            if key in item.text:
                item.text = item.text.replace(key, value)

def get_index(districtAlt,names):
    ind1=names.loc[names.str.contains("DISTRICT: {}".format(districtAlt))]
    
    start_ind=ind1.index[0]
    
    ind2=names.loc[names.str.contains('-')]
    indexTemp=np.where(ind2.index>start_ind)
    end_ind=ind2.index[indexTemp[0][0]]
    # print(ind1)
    return start_ind+1,end_ind
    
AWS_filename=filedialog.askopenfilename(title='Select the AWS file (CSV file)')
ARG_filename=filedialog.askopenfilename(title='Select ARG file')
# print(AWS_filename)
df=pd.read_csv(AWS_filename,skiprows=7,usecols=[1,2,5],header=None)
di=df[1]
stations=df[2]
rf=df[5]
df3=pd.read_csv(ARG_filename,skiprows=7,usecols=[1,2,5],header=None)
di3=df3[1]
stations3=df3[2]
rf3=df3[5]
time1=input('Enter Time:')
time2=input('Enter Time(-30m):')
uNo=input('Enter Update number:')
# now_utc = datetime(year=2021,month=1,day=1,minutes=time1)
Districts=['ALAPPUZHA','ERNAKULAM','IDUKKI','KANNUR','KASARAGOD','KOLLAM','KOTTAYAM','KOZHIKODE',\
            'MALAPPURAM','PALAKKAD','PATHANAMTHITTA','THIRUVANANTHAPURAM','THRISSUR',\
                'WAYANAD']
Districts_alt_name=['ALAPUZHA','ERNAKULAM','IDUKKI','CANNUR','KASARGOD','KOLLAM','KOTTAYAM','KOZHIKODE',\
                    'MALAPPURAM','PALAKKAD','PATHANAMTHITTA','THIRUVANANTHAPURAM','THRISSUR',\
                        'WYNAD']
if int(time1)<930:
    # dTemp1 = yeste/rday.strftime("%d/%m/%Y")
    index=int(yesterday.strftime("%d"))
else:
    # d1 = today.strftime("%d/%m/%Y")
    index=int(today.strftime("%d"))

d1 = today.strftime("%d/%m/%Y")
# now_asia = now_utc.astimezone(timezone('Asia/Kolkata')) 
df2=pd.read_excel("Rainfall Statement - MKFormat.xls",skiprows=10)
df2=df2.fillna('-')
names=df2["MET.SUB/DISTRICT/STATION"]
drms_rf=df2[index]
i_temp=0

dstrt = yesterday.strftime("%d/%m/%Y")



for district,districtAlt in zip(Districts,Districts_alt_name):
    start_ind,end_ind=get_index(districtAlt,names)
    stat=stations[di==district]
    stat3=stations3[di3==district]
    data=""
    # data3=""
    for station in stat:
        ind=rf[stations==station].index[0]
        rfVal=rf[stations==station]
        data+=("{} : {}\n".format(station, rfVal[ind]))
    # loc1=str(data)
    for station in stat3:
        ind=rf3[stations3==station].index[0]
        rfVal3=rf3[stations3==station]
        data+=("{} : {}\n".format(station, rfVal3[ind]))
    loc1=str(data)
    # print(loc3)
    # rain1="62.0"
    data2=""
    info="Rainfall received in mm during 24hrs ended at 0830 hrs IST of today"
    if int(time2)<830:
        info2="Rainfall received in mm at AWS stations from 0830 hrs IST of {} to {} hrs IST of {}.".format(dstrt,time2,d1)
    elif int(time2)>=2330 and int(time2)<2400:
        info2="Rainfall received in mm at AWS stations from 0830 hrs IST to {} hrs IST of {}.".format(time2,dstrt)
    else:
        info2="Rainfall received in mm at AWS stations from 0830 hrs IST  to {} hrs IST of Today.".format(time2)
    for indTemp in range(start_ind,end_ind):
        data2+=("{} : {}\n".format(names[indTemp],drms_rf[indTemp]))
    loc2=str(data2)
    loc2=loc2.rjust(25)
    i_temp+=1
    Caption=["Latest Satellite Picture","Latest Radar TVM","Latest Radar KOC"]
    variables = {
        "${DATE}": d1,
        "${TIME}": time1,
        "${Loc1}": loc2,
        # "${TIME2}":time2,
        "${DISTRICT}":district,
        "${UPDATENO}":uNo,
        "${Info}":info,
        "${Loc2}":loc1,
        "${Info2}":info2,
        # "${ARG}":data3,
        # "${Rain1}":rain1,
        # "${INFO}":info,
    }
    temlate_file_path = 'Template.docx'
    output_file_path = 'IBF\IBF_{}.docx'.format(district)
    template_document=Document(temlate_file_path)
    for variable_key, variable_value in variables.items():
        for paragraph in template_document.paragraphs:
            replace_text_in_paragraph(paragraph, variable_key, variable_value)
        for table in template_document.tables:
            for col in table.columns:
                for cell in col.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, variable_key, variable_value)
    
    tables = template_document.tables
    p = tables[2].rows[0].cells[0].add_paragraph()
    r = p.add_run()
    r.add_text("\n"+Caption[0])
    r.add_picture(Savefile[0],width=Inches(5), height=Inches(5))

    i=2 #ignoring trv radar image. To include, change i=1
    while i<3:
        p = tables[2].rows[i].cells[0].add_paragraph()
        r = p.add_run()
        r.add_text("\n"+Caption[i])
        r.add_picture(Savefile[i],width=Inches(7), height=Inches(5))

        i+=1
    template_document.save(output_file_path)
    if ibfdist[ibfyes] == 1:
        newname = 'IBF\{}\IBF_{}_{}_{}.docx'.format(district,district,today.strftime("%d%m%Y"),time1)
        template_document.save(newname)
    ibfyes = ibfyes + 1



    